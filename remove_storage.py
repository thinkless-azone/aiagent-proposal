from docx import Document

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        paragraph.text = paragraph.text.replace(old_text, new_text)
        print(f"Replaced in paragraph: {old_text} -> {new_text}")

def replace_text_in_table(table, old_text, new_text):
    for row in table.rows:
        for cell in row.cells:
            if old_text in cell.text:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, old_text, new_text)

# Load the v9 document
doc_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v9.docx"
doc = Document(doc_path)

# We need to remove the row with "СХД Raidix" or similar
# And update the total price.
# Previous total was ~32M (server) + ~9.8M (storage) + ~0.5M (switch) + software...
# We need to be careful with total price calculation if we don't have the exact number.
# But we can remove the row first.

rows_to_delete = []

for table in doc.tables:
    for row in table.rows:
        # Check if this row is for the storage
        is_storage_row = False
        for cell in row.cells:
            if "СХД" in cell.text and "Raidix" in cell.text:
                is_storage_row = True
                break
            if "Aerodisk" in cell.text: # Check for alternative name just in case
                 is_storage_row = True
                 break
        
        if is_storage_row:
            rows_to_delete.append(row)
            print("Found storage row to delete")

# Delete rows
for row in rows_to_delete:
    # This is a bit tricky in python-docx, usually we remove the element
    tbl = row._element.getparent()
    tbl.remove(row._element)
    print("Deleted storage row")

# Save as v10
output_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v10.docx"
doc.save(output_path)
print(f"Document saved to {output_path}")
