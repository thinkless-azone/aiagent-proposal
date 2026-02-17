from docx import Document

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        paragraph.text = paragraph.text.replace(old_text, new_text)
        print(f"Replaced in paragraph: {old_text} -> {new_text}")

def replace_text_in_table(table, old_text, new_text):
    for row in table.rows:
        for cell in row.cells:
            if old_text in cell.text:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, old_text, new_text)

# Load the v6 document
doc_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v6.docx"
doc = Document(doc_path)

replacements = [
    ("Гравитон С2122ИУ", "YADRO G4208P G3"),
    ("Graviton C2122IU", "YADRO G4208P G3"),
    ("3 750 000", "32 151 159"), # Replace price if present in text
]

# Replace in paragraphs
for paragraph in doc.paragraphs:
    for old, new in replacements:
        replace_text_in_paragraph(paragraph, old, new)

# Replace in tables
for table in doc.tables:
    for old, new in replacements:
        replace_text_in_table(table, old, new)

# Save as v7
output_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v7.docx"
doc.save(output_path)
print(f"Document saved to {output_path}")
