from docx import Document

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        paragraph.text = paragraph.text.replace(old_text, new_text)
        print(f"Replaced in paragraph: {old_text} -> {new_text}")

def replace_text_in_table(table, old_text, new_text):
    for row in table.rows:
        for cell in row.cells:
            if old_text in cell.text:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, old_text, new_text)

# Load the v8 document
doc_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v8.docx"
doc = Document(doc_path)

# We need to find the row with Eltex MES5324 and update quantity from 2 to 1
# and update the total price for that item (499,000 * 2 = 998,000 -> 499,000)

for table in doc.tables:
    for row in table.rows:
        # Check if this row is for the switch
        is_switch_row = False
        for cell in row.cells:
            if "Eltex MES5324" in cell.text:
                is_switch_row = True
                break
        
        if is_switch_row:
            # Try to find "2" in this row and replace with "1"
            for cell in row.cells:
                if cell.text.strip() == "2":
                    cell.text = "1"
                    print("Replaced quantity 2 -> 1 in switch row")
                # Update total price cell
                if "998 000" in cell.text:
                     cell.text = cell.text.replace("998 000", "499 000")
                     print("Replaced total price in switch row")

# Save as v9
output_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v9.docx"
doc.save(output_path)
print(f"Document saved to {output_path}")
