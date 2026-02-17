from docx import Document

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        inline = paragraph.runs
        # Simple replacement might break formatting if text is split across runs
        # But for simple model names it usually works or we can just replace the text of the paragraph
        # A safer approach for simple replacements:
        paragraph.text = paragraph.text.replace(old_text, new_text)
        print(f"Replaced in paragraph: {old_text} -> {new_text}")

def replace_text_in_table(table, old_text, new_text):
    for row in table.rows:
        for cell in row.cells:
            if old_text in cell.text:
                # Iterate over paragraphs in cell
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, old_text, new_text)

# Load the v5 document
doc_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v5.docx"
doc = Document(doc_path)

old_model = "MES2324"
new_model = "MES2300-24"

# Replace in paragraphs
for paragraph in doc.paragraphs:
    replace_text_in_paragraph(paragraph, old_model, new_model)

# Replace in tables
for table in doc.tables:
    replace_text_in_table(table, old_model, new_model)

# Save as v6
output_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v6.docx"
doc.save(output_path)
print(f"Document saved to {output_path}")
