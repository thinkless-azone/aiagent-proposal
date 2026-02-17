from docx import Document

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        paragraph.text = paragraph.text.replace(old_text, new_text)
        print(f"Replaced in paragraph: {old_text} -> {new_text}")

def replace_text_in_table(table, old_text, new_text):
    for row in table.rows:
        for cell in row.cells:
            if old_text in cell.text:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, old_text, new_text)

# Load the v7 document
doc_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v7.docx"
doc = Document(doc_path)

replacements = [
    ("2 шт", "1 шт"), # General replacement, might be risky, but context usually helps
    ("Количество: 2", "Количество: 1"),
    ("64 302 318", "32 151 159"), # Update total price for servers (2 -> 1)
]

# Specific logic for table cells might be needed if "2" is in a separate cell
# For now, let's try to find the row with YADRO and update the quantity cell if possible
# But since we don't have easy cell addressing, we'll rely on text replacement for now
# and manual verification if needed. 
# Given the previous replacements, the quantity might be in the text description or a separate column.

# Let's try to be more specific with replacements based on previous knowledge of the doc structure
# If the doc has a table with "YADRO G4208P G3" and a quantity column "2", we need to find it.

for table in doc.tables:
    for row in table.rows:
        # Check if this row is for the server
        is_server_row = False
        for cell in row.cells:
            if "YADRO G4208P G3" in cell.text:
                is_server_row = True
                break
        
        if is_server_row:
            # Try to find "2" in this row and replace with "1"
            for cell in row.cells:
                if cell.text.strip() == "2":
                    cell.text = "1"
                    print("Replaced quantity 2 -> 1 in server row")
                # Also update the total price cell if it exists and contains the old total
                if "64 302 318" in cell.text: # 32,151,159 * 2
                     cell.text = cell.text.replace("64 302 318", "32 151 159")
                     print("Replaced total price in server row")

# Save as v8
output_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v8.docx"
doc.save(output_path)
print(f"Document saved to {output_path}")
