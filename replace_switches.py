from docx import Document

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        paragraph.text = paragraph.text.replace(old_text, new_text)
        print(f"Replaced in paragraph: {old_text} -> {new_text}")

def replace_text_in_table(table, old_text, new_text):
    for row in table.rows:
        for cell in row.cells:
            if old_text in cell.text:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, old_text, new_text)

# Load the v10 document
doc_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v10.docx"
doc = Document(doc_path)

# 1. Replace MES2324 with MES2300-24 (if any left)
# 2. Replace MES5324 with MES2300-24
# 3. Update price for MES5324 (499,000) to MES2300-24 (139,000)

replacements = [
    ("MES2324", "MES2300-24"),
    ("MES5324", "MES2300-24"),
    ("499 000", "139 000"), # Price update for the switch row
]

for table in doc.tables:
    for row in table.rows:
        # Check if this row is for the switch
        is_switch_row = False
        for cell in row.cells:
            if "MES5324" in cell.text:
                is_switch_row = True
                # Replace text in this row
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, "MES5324", "MES2300-24")
            elif "MES2324" in cell.text:
                 # Replace text in this row
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, "MES2324", "MES2300-24")
        
        if is_switch_row:
            # Update price in this row
            for cell in row.cells:
                if "499 000" in cell.text:
                     cell.text = cell.text.replace("499 000", "139 000")
                     print("Replaced price 499 000 -> 139 000 in switch row")

# Also do general text replacement for descriptions outside tables
for paragraph in doc.paragraphs:
    replace_text_in_paragraph(paragraph, "MES2324", "MES2300-24")
    replace_text_in_paragraph(paragraph, "MES5324", "MES2300-24")

# Save as v11
output_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v11.docx"
doc.save(output_path)
print(f"Document saved to {output_path}")
