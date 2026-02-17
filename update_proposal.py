from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)  # Dark Blue
    p.space_before = Pt(18)
    p.space_after = Pt(12)

def add_subsection_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(51, 51, 51)  # Dark Gray
    p.space_before = Pt(12)
    p.space_after = Pt(6)

def add_bullet_point(doc, title, description):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    run_bullet = p.add_run("• ")
    run_title = p.add_run(title + ": ")
    run_title.bold = True
    p.add_run(description)

# Load the document
doc_path = "/home/ubuntu/upload/Проектноепредложение3.docx"
doc = Document(doc_path)

# Add new section
add_section_header(doc, "ДОПОЛНИТЕЛЬНЫЕ СЦЕНАРИИ ИСПОЛЬЗОВАНИЯ")

p = doc.add_paragraph("Помимо основной функции обеспечения экономической безопасности, система предоставляет широкие возможности для автоматизации финансовых процессов и поддержки принятия управленческих решений.")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Financial Operations
add_subsection_header(doc, "1. Финансовые операции и Закупки")
add_bullet_point(doc, "Автоматизация учёта расходов", "Распознавание и классификация счетов-фактур, автоматический ввод данных в 1С с распределением по статьям затрат.")
add_bullet_point(doc, "Контроль кассовых разрывов", "Прогнозирование дефицита денежных средств на основе анализа платёжного календаря и исторических данных.")
add_bullet_point(doc, "Анализ поставщиков", "Глубокая проверка контрагентов, анализ истории цен и надежности поставок.")
add_bullet_point(doc, "Анализ контрактов", "Автоматическая проверка договоров на соответствие корпоративным стандартам и выявление рисков.")

# Management Decisions
add_subsection_header(doc, "2. Управленческие решения")
add_bullet_point(doc, "Дашборд финансового состояния", "Визуализация ключевых метрик (маржинальность, ROI, коэффициенты ликвидности) в режиме реального времени.")
add_bullet_point(doc, "Сценарное моделирование", "Расчёт влияния изменения цен, объёмов продаж или курсов валют на финансовый результат компании.")
add_bullet_point(doc, "Анализ рентабельности по клиентам", "Выявление убыточных и высокомаржинальных клиентов для оптимизации клиентского портфеля.")
add_bullet_point(doc, "Прогнозирование спроса", "Использование ML-моделей для планирования производства и управления товарными запасами.")

# Construction & PTO
add_subsection_header(doc, "3. ПТО и Строительство")
add_bullet_point(doc, "Автоматизация смет", "Интеллектуальный анализ проектной документации и автоматическое формирование сметных расчетов.")
add_bullet_point(doc, "Система документооборота", "Автоматическая классификация, маршрутизация и контроль исполнительной документации.")
add_bullet_point(doc, "Видеоаналитика строительства", "Мониторинг прогресса работ, контроль соблюдения техники безопасности и использования ресурсов.")
add_bullet_point(doc, "Прогнозирование задержек", "ML-анализ графиков работ, поставок материалов и погодных условий для раннего выявления рисков срыва сроков.")

# Strategic Planning
add_subsection_header(doc, "4. Стратегическое планирование")
add_bullet_point(doc, "Бюджетирование на основе данных", "Анализ исторических тенденций для более точного и обоснованного планирования расходов.")
add_bullet_point(doc, "Оценка инвестиционных проектов", "Автоматический расчёт NPV, IRR и других показателей эффективности для принятия решений о капиталовложениях.")
add_bullet_point(doc, "Анализ конкурентоспособности", "Сравнение цен и условий с рыночными показателями для корректировки стратегии позиционирования.")

# Save the updated document
output_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v3.docx"
doc.save(output_path)
print(f"Document saved to {output_path}")
