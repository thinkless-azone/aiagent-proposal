import os
from docx import Document
from docx.shared import Pt

def update_implementation_plan(doc_path, output_path):
    doc = Document(doc_path)
    
    # Define replacements for implementation plan
    replacements = {
        "Внедрение базового варианта - за 4 месяца": "Внедрение базового варианта - за 3 месяца",
        "Внедрения оптимального варианта - от 6 месяцев": "Внедрения оптимального варианта - от 3 до 6 месяцев",
        "развёртывание кластера из 2-х серверов Гравитон": "развёртывание сервера YADRO G4208P",
        "Монтаж кластера (2x Гравитон С2122ИУ)": "Монтаж сервера YADRO G4208P",
        "Проектирование HA-кластера": "Проектирование архитектуры (YADRO G4208P)"
    }

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
                # Re-apply style if needed (basic check)
                if "Heading" in paragraph.style.name:
                    pass 

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)

    doc.save(output_path)
    print(f"Updated proposal saved to {output_path}")

if __name__ == "__main__":
    # Find the latest proposal file
    files = [f for f in os.listdir(".") if f.startswith("Commercial_Proposal_v") and f.endswith(".docx")]
    if not files:
        print("No proposal file found.")
        exit(1)
    
    latest_file = sorted(files, key=lambda x: int(x.split('_v')[1].split('.')[0]))[-1]
    version_num = int(latest_file.split('_v')[1].split('.')[0])
    new_version_num = version_num + 1
    output_file = f"Commercial_Proposal_v{new_version_num}.docx"
    
    print(f"Updating {latest_file} -> {output_file}")
    update_implementation_plan(latest_file, output_file)
