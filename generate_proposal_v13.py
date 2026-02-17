from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def generate_proposal():
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title Page
    doc.add_paragraph('\n\n\n\n\n')
    title = doc.add_paragraph('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(24)
    title.style.font.bold = True
    title.style.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph('Внедрение локального ИИ-агента\nдля автоматизации 1С:Предприятие')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style.font.size = Pt(16)
    subtitle.style.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph('\n\n\n')
    
    info = doc.add_paragraph(f'Дата: {datetime.datetime.now().strftime("%d.%m.%Y")}')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n\n\n\n')
    
    footer = doc.add_paragraph('Конфиденциально. Только для внутреннего пользования.')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.style.font.size = Pt(9)
    footer.style.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # 1. Executive Summary
    h1 = doc.add_heading('1. Резюме проекта', level=1)
    h1.style.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph(
        'Предлагаемое решение представляет собой программно-аппаратный комплекс для внедрения '
        'автономного ИИ-агента в инфраструктуру заказчика. Система обеспечивает автоматизацию '
        'рутинных операций в 1С, проверку контрагентов и анализ договоров без передачи данных '
        'во внешние облачные сервисы.'
    )

    # 2. Proposed Solution (Optimal Only)
    h1 = doc.add_heading('2. Предлагаемое решение', level=1)
    h1.style.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph(
        'Мы предлагаем оптимальную конфигурацию на базе высокопроизводительного сервера YADRO, '
        'специализированного для задач искусственного интеллекта. Данное решение обеспечивает '
        'максимальную производительность, масштабируемость и надежность.'
    )

    # Hardware Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False
    
    # Set column widths
    for cell in table.columns[0].cells: cell.width = Inches(0.5)
    for cell in table.columns[1].cells: cell.width = Inches(3.0)
    for cell in table.columns[2].cells: cell.width = Inches(0.8)
    for cell in table.columns[3].cells: cell.width = Inches(1.5)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Наименование'
    hdr_cells[2].text = 'Кол-во'
    hdr_cells[3].text = 'Цена (₽)'

    # Make header bold and gray background
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E6E6E6')
        tcPr.append(shd)

    items = [
        ('Сервер YADRO G4208P G3 (2x Intel Xeon 6526Y, 512GB RAM, 1x NVIDIA H100 80GB, 2x 960GB SSD, 12x 3.5" HDD Slots)', 1, 32151159),
        ('Коммутатор Eltex MES2300-24 (24x 1G, 4x 10G SFP+)', 1, 139000),
        ('Лицензия "ИИ-Агент: Корпоративный" (бессрочная)', 1, 1500000),
        ('Модуль интеграции 1С (расширение конфигурации)', 1, 450000),
        ('Пакет внедрения и настройки (под ключ)', 1, 1200000)
    ]

    total = 0
    for i, (name, qty, price) in enumerate(items, 1):
        row_cells = doc.add_table(rows=1, cols=4).rows[0].cells # Hack to get new row cells
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = name
        row_cells[2].text = str(qty)
        row_cells[3].text = f'{price:,.0f}'.replace(',', ' ')
        total += price

    # Total Row
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[2])
    row_cells[0].text = 'ИТОГО:'
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    row_cells[3].text = f'{total:,.0f}'.replace(',', ' ')
    row_cells[3].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph('\n')
    doc.add_paragraph(f'Общая стоимость проекта: {total:,.0f} рублей (без НДС).'.replace(',', ' '))

    # 3. Implementation Plan
    h1 = doc.add_heading('3. План внедрения', level=1)
    h1.style.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph('Срок реализации проекта: 3 - 6 месяцев.')

    plan_items = [
        ('Этап 1: Глубокий анализ (1 месяц)', 'Аудит бизнес-процессов, проектирование архитектуры, детальное ТЗ.'),
        ('Этап 2: Развертывание ядра (2-3 месяца)', 'Монтаж сервера YADRO, настройка среды ИИ, тюнинг моделей.'),
        ('Этап 3: Внедрение (1-2 месяца)', 'Нагрузочное тестирование, обучение персонала, запуск в эксплуатацию.')
    ]

    for title, desc in plan_items:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.add_run(f'\n{desc}')

    # 4. Benefits
    h1 = doc.add_heading('4. Ожидаемые результаты', level=1)
    h1.style.font.color.rgb = RGBColor(0, 51, 102)

    benefits = [
        'Полная автономность и безопасность данных (On-Premise).',
        'Высокая производительность благодаря ускорителю NVIDIA H100.',
        'Масштабируемость системы за счет внутренней емкости сервера YADRO.',
        'Импортозамещение: использование оборудования из реестра Минпромторга.'
    ]

    for benefit in benefits:
        doc.add_paragraph(benefit, style='List Bullet')

    # Footer with page numbers
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Страница "
    add_page_number(p.add_run())

    doc.save('Commercial_Proposal_v13.docx')
    print("Proposal v13 generated successfully.")

if __name__ == "__main__":
    generate_proposal()
