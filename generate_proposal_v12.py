from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.text = "Страница "
    add_page_number(run)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.space_after = Pt(24)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)
        p.space_before = Pt(18)
        p.space_after = Pt(12)
    else:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(51, 51, 51)
        p.space_before = Pt(12)
        p.space_after = Pt(6)

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.space_after = Pt(8)

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Pt(18)

def create_proposal():
    doc = Document()
    set_footer(doc)

    # Title Page
    add_title(doc, "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ")
    p = doc.add_paragraph("Локальный ИИ-агент для 1С:Предприятие")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(18)
    
    doc.add_page_break()

    # 1. Project Overview
    add_heading(doc, "1. ОПИСАНИЕ ПРОЕКТА")
    add_paragraph(doc, "Предлагаемое решение представляет собой программно-аппаратный комплекс для автоматизации финансово-хозяйственных операций с использованием технологий искусственного интеллекта. Система разворачивается полностью локально (On-Premise), обеспечивая максимальную безопасность данных и соответствие требованиям импортозамещения.")

    # 2. Hardware Specification
    add_heading(doc, "2. СПЕЦИФИКАЦИЯ ОБОРУДОВАНИЯ")
    
    # Optimal Variant
    add_heading(doc, "Оптимальный вариант (Рекомендуемый)", level=2)
    add_paragraph(doc, "Высокопроизводительная конфигурация на базе сервера YADRO для задач обучения и инференса ИИ.")
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Наименование'
    hdr_cells[1].text = 'Кол-во'
    hdr_cells[2].text = 'Цена за ед. (₽)'
    hdr_cells[3].text = 'Сумма (₽)'

    items = [
        ("Сервер YADRO G4208P G3 (2x Intel Xeon 6526Y, 512GB RAM, 1x NVIDIA H100 80GB, 2x 960GB SSD, 12x 3.5\" LFF)", 1, 32151159, 32151159),
        ("Коммутатор Eltex MES2300-24 (24x 1G, 4x 10G SFP+)", 1, 139000, 139000),
        ("ИБП Ippon Innova RT 3000 (3000VA/2700W)", 1, 120000, 120000),
        ("Шкаф телекоммуникационный 42U (ЦМО)", 1, 85000, 85000),
        ("Комплект кабелей и монтажных материалов", 1, 50000, 50000)
    ]

    total = 0
    for item, qty, price, subtotal in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = str(qty)
        row_cells[2].text = f"{price:,.0f}".replace(",", " ")
        row_cells[3].text = f"{subtotal:,.0f}".replace(",", " ")
        total += subtotal

    add_paragraph(doc, f"\nИтого оборудование: {total:,.0f} ₽".replace(",", " "))

    # Basic Variant
    add_heading(doc, "Базовый вариант", level=2)
    add_paragraph(doc, "Экономичное решение для начального этапа внедрения.")
    
    table_basic = doc.add_table(rows=1, cols=4)
    table_basic.style = 'Table Grid'
    hdr_cells = table_basic.rows[0].cells
    hdr_cells[0].text = 'Наименование'
    hdr_cells[1].text = 'Кол-во'
    hdr_cells[2].text = 'Цена за ед. (₽)'
    hdr_cells[3].text = 'Сумма (₽)'

    items_basic = [
        ("Сервер Гравитон С2101И (2x Intel Xeon 4310, 128GB RAM, 2x 960GB SSD)", 1, 1200000, 1200000),
        ("Коммутатор Eltex MES2300-24 (24x 1G, 4x 10G SFP+)", 1, 139000, 139000),
        ("ИБП Ippon Smart Winner 2000 (2000VA)", 1, 65000, 65000),
        ("Шкаф настенный 12U", 1, 25000, 25000)
    ]

    total_basic = 0
    for item, qty, price, subtotal in items_basic:
        row_cells = table_basic.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = str(qty)
        row_cells[2].text = f"{price:,.0f}".replace(",", " ")
        row_cells[3].text = f"{subtotal:,.0f}".replace(",", " ")
        total_basic += subtotal

    add_paragraph(doc, f"\nИтого оборудование: {total_basic:,.0f} ₽".replace(",", " "))

    # 3. Implementation Plan
    add_heading(doc, "3. ПЛАН ВНЕДРЕНИЯ")
    
    add_heading(doc, "Базовый вариант (3 месяца)", level=2)
    add_bullet(doc, "Месяц 1: Анализ инфраструктуры, закупка оборудования (Гравитон С2101И), разработка ТЗ.")
    add_bullet(doc, "Месяц 2-2.5: Монтаж сервера, установка ПО (Альт Линукс, Postgres Pro), базовая настройка моделей.")
    add_bullet(doc, "Месяц 3: Обучение персонала, опытная эксплуатация, передача в поддержку.")

    add_heading(doc, "Оптимальный вариант (3-6 месяцев)", level=2)
    add_bullet(doc, "Месяц 1: Глубокий аудит бизнес-процессов, проектирование архитектуры (YADRO G4208P), детальное ТЗ.")
    add_bullet(doc, "Месяц 2-4: Монтаж сервера YADRO G4208P, настройка высокопроизводительной среды ИИ, тюнинг моделей.")
    add_bullet(doc, "Месяц 5-6: Нагрузочное тестирование, обучение администраторов, запуск в промышленную эксплуатацию.")

    # 4. Additional Scenarios
    add_heading(doc, "4. ДОПОЛНИТЕЛЬНЫЕ СЦЕНАРИИ ИСПОЛЬЗОВАНИЯ")
    add_paragraph(doc, "Помимо основной функции обеспечения экономической безопасности, система предоставляет широкие возможности для автоматизации финансовых процессов и поддержки принятия управленческих решений.")
    
    add_heading(doc, "Финансовые операции и Закупки", level=2)
    add_bullet(doc, "Автоматизация учёта расходов: Распознавание и классификация счетов-фактур.")
    add_bullet(doc, "Контроль кассовых разрывов: Прогнозирование дефицита денежных средств.")
    add_bullet(doc, "Анализ поставщиков: Глубокая проверка контрагентов.")

    add_heading(doc, "ПТО и Строительство", level=2)
    add_bullet(doc, "Автоматизация смет: Интеллектуальный анализ проектной документации.")
    add_bullet(doc, "Видеоаналитика строительства: Мониторинг прогресса работ.")
    add_bullet(doc, "Прогнозирование задержек: ML-анализ графиков работ.")

    # Save
    output_path = "/home/ubuntu/ai-agent-proposal/client/public/Commercial_Proposal_v12.docx"
    doc.save(output_path)
    print(f"Proposal v12 generated at: {output_path}")

if __name__ == "__main__":
    create_proposal()
