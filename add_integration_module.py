from docx import Document
from docx.shared import Pt, RGBColor

def add_section_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)  # Dark Blue
    p.space_before = Pt(18)
    p.space_after = Pt(12)

def add_subsection_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(51, 51, 51)  # Dark Gray
    p.space_before = Pt(12)
    p.space_after = Pt(6)

def add_bullet_point(doc, title, description):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    run_bullet = p.add_run("• ")
    run_title = p.add_run(title + ": ")
    run_title.bold = True
    p.add_run(description)

# Load the v4 document
doc_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v4.docx"
doc = Document(doc_path)

# Add new section for Integration Module
add_section_header(doc, "МОДУЛЬ ИНТЕГРАЦИИ С 1С:ПРЕДПРИЯТИЕ")

p = doc.add_paragraph("Система включает специализированный модуль для бесшовной интеграции с существующей конфигурацией 1С заказчика, обеспечивающий двусторонний обмен данными в реальном времени.")

# Architecture
add_subsection_header(doc, "1. Архитектура и Протоколы")
add_bullet_point(doc, "REST API", "Использование стандартного интерфейса OData и HTTP-сервисов платформы 1С для универсального доступа к данным.")
add_bullet_point(doc, "Web-сервисы (SOAP)", "Поддержка классических SOAP-протоколов для интеграции с устаревшими конфигурациями.")
add_bullet_point(doc, "Очереди сообщений", "Асинхронный обмен через RabbitMQ/Kafka для высоконагруженных систем, гарантирующий доставку сообщений.")

# Data Formats
add_subsection_header(doc, "2. Форматы обмена данными")
add_bullet_point(doc, "JSON", "Основной формат для передачи легковесных структур данных и команд управления.")
add_bullet_point(doc, "XML / XDTO", "Строгая типизация данных для сложных документов и справочников, соответствующая стандартам Enterprise Data.")
add_bullet_point(doc, "CSV / Excel", "Поддержка пакетной загрузки и выгрузки исторических данных для первоначального обучения моделей.")

# Security
add_subsection_header(doc, "3. Безопасность и Контроль")
add_bullet_point(doc, "Аутентификация", "Поддержка OAuth 2.0 и Basic Auth с использованием служебных пользователей 1С с ограниченными правами.")
add_bullet_point(doc, "Шифрование", "Весь трафик передается по защищенному протоколу HTTPS (TLS 1.3).")
add_bullet_point(doc, "Журналирование", "Полный аудит всех запросов и изменений данных в журнале регистрации 1С.")

# Save as v5
output_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v5.docx"
doc.save(output_path)
print(f"Document saved to {output_path}")
