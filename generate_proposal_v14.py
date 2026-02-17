
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

def create_proposal():
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title Page
    doc.add_paragraph('\n\n\n\n')
    title = doc.add_paragraph('КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(24)
    title.style.font.bold = True
    title.style.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Внедрение нейросетей в бизнес-процессы\nи создание локального ИИ-контура')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style.font.size = Pt(16)
    subtitle.style.font.color.rgb = RGBColor(80, 80, 80)
    
    doc.add_paragraph('\n\n\n\n\n\n')
    
    info = doc.add_paragraph(f'Дата: {datetime.datetime.now().strftime("%d.%m.%Y")}\nВерсия: 14.0 (Расширенная)\nДля: Руководства компании')
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
    # 1. Executive Summary
    h1 = doc.add_heading('1. Резюме проекта', level=1)
    h1.style.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph(
        'Предлагается комплексное решение по внедрению технологий искусственного интеллекта (ИИ) '
        'в ключевые бизнес-процессы компании. Проект выходит за рамки создания ИИ-агента для 1С '
        'и охватывает автоматизацию документооборота, закупок, контроля строительства и безопасности.'
    )
    
    doc.add_paragraph(
        'Решение строится на базе высокопроизводительного российского сервера YADRO G4208P G3 '
        'с ускорителем NVIDIA H100, что обеспечивает полную автономность, безопасность данных '
        'и независимость от облачных сервисов.'
    )
    
    # 2. Key AI Modules
    h1 = doc.add_heading('2. Ключевые направления внедрения ИИ', level=1)
    h1.style.font.color.rgb = RGBColor(0, 51, 102)
    
    modules = [
        {
            'title': 'Управление документацией (OCR + NLP)',
            'desc': 'Автоматическое распознавание и классификация входящих документов (акты, счета, договоры). '
                    'Умный семантический поиск по архиву за секунды. Версионирование и контроль изменений.'
        },
        {
            'title': 'Автоматизация смет и закупок',
            'desc': 'Генерация смет на основе текстовых описаний работ. Автоматическая сверка цен поставщиков '
                    'с рыночными индикаторами. Выявление аномалий в закупках и предотвращение переплат.'
        },
        {
            'title': 'Видеоаналитика и безопасность',
            'desc': 'Мониторинг соблюдения техники безопасности на объектах (каски, жилеты). '
                    'Контроль хода строительства и сравнение с планом-графиком. Антифрод-система для финансовых операций.'
        },
        {
            'title': 'Интеллектуальный ассистент 1С',
            'desc': 'Голосовой и текстовый интерфейс для работы с 1С:Предприятие. '
                    'Формирование отчетов по запросу ("Покажи продажи за прошлый месяц"), помощь в заполнении форм.'
        }
    ]
    
    for mod in modules:
        p = doc.add_paragraph()
        runner = p.add_run(mod['title'])
        runner.bold = True
        runner.font.color.rgb = RGBColor(0, 51, 102)
        p.add_run('\n' + mod['desc'])
    
    # 3. Technical Solution (Optimal Variant)
    h1 = doc.add_heading('3. Техническое решение (Оптимальный вариант)', level=1)
    h1.style.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph('Единая платформа для всех ИИ-сервисов:')
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Компонент'
    hdr_cells[1].text = 'Описание'
    
    specs = [
        ('Сервер', 'YADRO G4208P G3 (2x Intel Xeon Gold 6526Y, 512GB RAM)'),
        ('Ускоритель ИИ', '1x NVIDIA H100 80GB (Специализирован для LLM и Vision)'),
        ('Хранение данных', 'Внутреннее: 4x 3.84TB NVMe SSD (Быстрый доступ) + 12 слотов расширения'),
        ('Сеть', 'Коммутатор Eltex MES2300-24 (24x 1G, 4x 10G SFP+)'),
        ('ПО', 'ОС Альт Линукс СПТ, СУБД Postgres Pro Enterprise, Платформа ML')
    ]
    
    for item, desc in specs:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = desc
        
    doc.add_paragraph('\nПримечание: Внешняя СХД исключена, так как сервер обладает достаточной внутренней емкостью.')

    # 4. Implementation Plan
    h1 = doc.add_heading('4. План внедрения', level=1)
    h1.style.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph('Срок реализации: 3-6 месяцев')
    
    plan_steps = [
        'Этап 1 (Месяц 1): Поставка оборудования, развертывание инфраструктуры, установка базового ПО.',
        'Этап 2 (Месяц 2-3): Настройка моделей OCR и NLP, интеграция с 1С, пилотный запуск модуля документооборота.',
        'Этап 3 (Месяц 4-6): Обучение моделей видеоаналитики, запуск модуля закупок, масштабирование на все отделы.'
    ]
    
    for step in plan_steps:
        doc.add_paragraph(step, style='List Bullet')

    # 5. Budget Estimate
    h1 = doc.add_heading('5. Бюджетная оценка', level=1)
    h1.style.font.color.rgb = RGBColor(0, 51, 102)
    
    budget_table = doc.add_table(rows=1, cols=3)
    budget_table.style = 'Table Grid'
    hdr_cells = budget_table.rows[0].cells
    hdr_cells[0].text = 'Наименование'
    hdr_cells[1].text = 'Кол-во'
    hdr_cells[2].text = 'Стоимость (руб.)'
    
    items = [
        ('Сервер YADRO G4208P G3 (с NVIDIA H100)', 1, 32151159),
        ('Коммутатор Eltex MES2300-24', 1, 139000),
        ('Лицензии ПО (ОС, Виртуализация, СУБД)', 1, 3500000),
        ('Бизнес-модули ИИ (OCR, Сметы, Видео)', 1, 2450000),
        ('Работы по внедрению и настройке', 1, 4500000)
    ]
    
    total = 0
    for name, qty, price in items:
        row_cells = budget_table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = str(qty)
        row_cells[2].text = f'{price:,.0f}'.replace(',', ' ')
        total += price
        
    row_cells = budget_table.add_row().cells
    row_cells[0].text = 'ИТОГО'
    row_cells[2].text = f'{total:,.0f}'.replace(',', ' ')
    row_cells[0].paragraphs[0].runs[0].bold = True
    row_cells[2].paragraphs[0].runs[0].bold = True

    doc.save('Commercial_Proposal_v14.docx')
    print(f"Proposal v14 generated successfully. Total: {total:,.0f} RUB")

if __name__ == '__main__':
    create_proposal()
