from docx import Document

doc_path = "/home/ubuntu/ai-agent-proposal/client/public/Проектноепредложение3_v2.docx"
try:
    doc = Document(doc_path)
    print(f"Successfully opened {doc_path}")
    
    found = False
    for p in doc.paragraphs:
        if "ДОПОЛНИТЕЛЬНЫЕ СЦЕНАРИИ ИСПОЛЬЗОВАНИЯ" in p.text:
            found = True
            print("Found section header!")
            break
            
    if found:
        print("Verification SUCCESS: New section exists.")
        # Print a few lines after the header to confirm content
        printing = False
        count = 0
        for p in doc.paragraphs:
            if "ДОПОЛНИТЕЛЬНЫЕ СЦЕНАРИИ ИСПОЛЬЗОВАНИЯ" in p.text:
                printing = True
            if printing and count < 5:
                print(f"Content: {p.text[:50]}...")
                count += 1
    else:
        print("Verification FAILED: New section NOT found.")
        
except Exception as e:
    print(f"Error reading file: {e}")
